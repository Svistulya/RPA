from datetime import datetime
import pandas as pd
from docx import Document
from datetime import datetime


'''
    Вариант Г: Ввести данные вручную через консоль (минимум 5 строк).
    Вариант A: Рассчитать итоговую сумму для колонки "Стоимость" (Количество * Цена) и вставить итог (SUM) в отдельную ячейку.
    Вариант Б: Создать документ с заголовком, маркированным списком основных позиций и итоговой суммой.
    Вариант В: Создать PDF, вставив в него изображение (например, логотип) и итоговую текстовую сводку.
    Вариант Г: Упаковать все созданные файлы в ZIP-архив с именем "Отчет_[дата]".
'''

print("Введите данные")
data = []
count = 0
while True:
    name = input("Название: ").strip()
    qty = int(input("Количество: "))
    price = float(input("Цена за единицу: "))
    count += 1
    data.append({"Название": name, "Количество": qty, "Цена": price})

    if count < 5:
        continue
    else:
        break
    
print("\nВведённые данные:")
for i, item in enumerate(data, 1):
    print(f"{i}. {item['Название']} — {item['Количество']} шт. x {item['Цена']} руб.")

total_sum = 0
for item in data:
    cost = item["Количество"] * item["Цена"]
    item["Стоимость"] = cost
    total_sum += cost
    print(f"{item['Название']}: {item['Количество']} x {item['Цена']} = {cost:.2f} руб.")



df = pd.DataFrame(data)

df["Стоимость"] = df["Количество"] * df["Цена"]


total_cost = df["Стоимость"].sum()


total_row = {
    "Товар": "ИТОГО",
    "Количество": "",
    "Цена": "",
    "Стоимость": total_cost,
}

df_with_total = df._append(total_row, ignore_index=True)

file_name = "таблица.xlsx"
df_with_total.to_excel(file_name, index=False, sheet_name="Отчёт")

#Ворд

doc = Document()


title = doc.add_heading(f'Отчёт о покупках ')


doc.add_paragraph("Основные позиции:", style='Heading 2')


for item in data:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f"{item['Название']} — {item['Стоимость']:.2f} руб. "
              f"({item['Количество']} шт. по {item['Цена']} руб.)")


doc.add_paragraph()  
summary = doc.add_paragraph()
summary.add_run("Итого к оплате: ") 
summary.add_run(f"{total_sum:.2f} руб.") 


filename = "report.docx"
doc.save(filename)


#PDF

from fpdf import FPDF


pdf = FPDF()
pdf.add_page()
pdf.set_auto_page_break(auto=True, margin=15)


pdf.add_font("ArialUnicodeMS", style="", fname="") 
pdf.set_font("ArialUnicodeMS", size=12)


pdf.cell(0, 10, f"Итоговая сводка от {datetime.now().strftime('%d.%m.%Y')}", ln=True, align='C')


pdf.ln(10)

pdf.cell(0, 10, "Позиции в отчёте:", ln=True)

for item in data:
    pdf.cell(0, 8, f"• {item['Название']} — {item['Стоимость']:.2f} руб. "
                   f"({item['Количе12' \
                   'ство']} шт. по {item['Цена']} руб.)", ln=True)


pdf.ln(10)
pdf.set_font("ArialUnicodeMS", style="B", size=12)
pdf.cell(0, 10, f"ИТОГО К ОПЛАТЕ: {total_sum:.2f} руб.", ln=True, align='R')


filename = "summary.pdf"
pdf.output(filename)
